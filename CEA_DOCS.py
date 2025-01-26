import os
import logging
from pathlib import Path
from dotenv import load_dotenv
from docx import Document
from composio_crewai import ComposioToolSet, Action
from crewai import Crew, Agent, Task, Process
from langchain_openai import ChatOpenAI

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

# Initialize the language model
llm = ChatOpenAI(model="gpt-4o")

# Define paths
INPUT_DOC_PATH = Path("./Amarin_CEA_Design_Specification.docx")
OUTPUT_DIR = Path("./output")
OUTPUT_DIR.mkdir(exist_ok=True)
SUMMARY_DOC_PATH = OUTPUT_DIR / "summary.docx"

# Validate input file
if not INPUT_DOC_PATH.exists():
    raise FileNotFoundError(f"The document at {INPUT_DOC_PATH} does not exist.")

# Extract content from the Word document
def extract_docx_content(file_path):
    try:
        document = Document(file_path)
        content = [
            paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()
        ]
        if not content:
            logger.warning("The document is empty.")
        else:
            logger.info("Document content extracted successfully.")
        return content
    except Exception as e:
        logger.error(f"Error reading the Word document: {e}")
        raise

# Save summary to a Word document
def save_summary_to_docx(content, output_path):
    try:
        doc = Document()
        doc.add_heading("Document Summary", level=1)
        if isinstance(content, str):
            content = [content]
        for paragraph in content:
            doc.add_paragraph(paragraph)
        doc.save(output_path)
        logger.info(f"Summary saved to {output_path}")
    except Exception as e:
        logger.error(f"Error saving summary to .docx: {e}")
        raise

# Dynamic Prompt for Agent Execution
def create_dynamic_prompt(content):
    return (
        "You are an expert summarization assistant. Your goal is to analyze the provided content, distill the main points, "
        "and generate a concise, structured summary. Focus on readability and simplifying complex ideas for a professional audience. "
        f"Content to analyze:\n{content}\n\n"
        "Output the summary in this format:\n"
        "- Main Topics: [List key themes]\n"
        "- Summary: [Provide a concise summary]\n"
        "- Insights: [Include actionable insights, if applicable]"
    )

# Define the Composio ToolSet
composio_toolset = ComposioToolSet(output_dir=OUTPUT_DIR)
tools = composio_toolset.get_tools(actions=[
    Action.CODEINTERPRETER_EXECUTE_CODE,
    Action.CODEINTERPRETER_GET_FILE_CMD,
    Action.CODEINTERPRETER_RUN_TERMINAL_CMD,
])

# Define the Agent
ppt_creator_agent = Agent(
    role="Content Analysis Expert",
    goal="Analyze the provided document and create a professional summary with clear structure.",
    backstory=(
        "You are an AI assistant trained to analyze complex documents and produce professional-grade summaries. "
        "Your output must be actionable and readable for individuals with varying levels of expertise."
    ),
    tools=tools,
)

# Define the Task
summary_task = Task(
    description="Analyze the document content and generate a structured summary with key themes, concise explanations, and actionable insights.",
    expected_output="A professional summary saved as a .docx file.",
    tools=tools,
    agent=ppt_creator_agent,
    verbose=True,
)

# Run the process
def run_process():
    try:
        # Extract document content
        content = extract_docx_content(INPUT_DOC_PATH)

        if not content:
            logger.error("No content to process. Ensure the document has readable text.")
            return

        # Create a dynamic prompt
        dynamic_prompt = create_dynamic_prompt(content)
        logger.info(f"Dynamic Prompt for Task: {dynamic_prompt}")

        # Update task description dynamically
        summary_task.description = dynamic_prompt

        # Create a Crew for sequential task execution
        crew = Crew(agents=[ppt_creator_agent], tasks=[summary_task], process=Process.sequential)

        # Kickoff the process
        creation_response = crew.kickoff()

        # Debugging: Log the raw response structure
        logger.info(f"Raw Creation Response: {creation_response}")

        # Access the task output (adjust based on actual structure)
        if hasattr(creation_response, 'tasks') and creation_response.tasks:
            task_output = creation_response.tasks[0].output  # Assuming tasks is a list
        else:
            logger.error("No task output received from the agent.")
            return

        # Save summarized content to a Word document
        save_summary_to_docx(task_output, SUMMARY_DOC_PATH)

        logger.info(f"Process completed successfully. Summary available at {SUMMARY_DOC_PATH}")
    except Exception as e:
        logger.error(f"Process failed: {e}")
        raise




# Main execution
if __name__ == "__main__":
    run_process()
