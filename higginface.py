from transformers import AutoTokenizer, pipeline
import torch
from docx import Document

# Initialize summarization pipeline and tokenizer
model_name = "facebook/bart-large-cnn"
summarizer = pipeline(task="summarization", model=model_name, device=0 if torch.cuda.is_available() else -1)
tokenizer = AutoTokenizer.from_pretrained(model_name)

# Function to extract text from .docx
def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        text = " ".join(paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip())
        return text
    except Exception as e:
        print(f"Error reading the .docx file: {e}")
        return ""

# Function to split text into manageable chunks
def split_text(text, tokenizer, max_tokens=1024):
    words = text.split()
    chunks = []
    current_chunk = []

    for word in words:
        current_chunk.append(word)
        # Check token length of the current chunk
        if len(tokenizer(" ".join(current_chunk))["input_ids"]) >= max_tokens:
            chunks.append(" ".join(current_chunk[:-1]))  # Add all but the last word
            current_chunk = [word]  # Start the next chunk with the last word

    if current_chunk:
        chunks.append(" ".join(current_chunk))

    return chunks

# Function to summarize text chunks
def summarize_text(text, tokenizer, chunk_size=1024):
    chunks = split_text(text, tokenizer, max_tokens=chunk_size)
    summaries = []

    for i, chunk in enumerate(chunks):
        try:
            summary = summarizer(chunk, max_length=200, min_length=50, do_sample=False)
            summaries.append(summary[0]["summary_text"])
        except Exception as e:
            print(f"Error summarizing chunk {i}: {e}")

    return " ".join(summaries)

# Function to save summary to a Word document
def save_summary_to_doc(summary, output_path):
    try:
        # Create a new Word document
        doc = Document()
        doc.add_heading("Summarized Content", level=1)  # Add a heading
        doc.add_paragraph(summary)  # Add the summary text
        doc.save(output_path)  # Save the document to the specified path
        print(f"Summary saved successfully to {output_path}")
    except Exception as e:
        print(f"Error saving summary to Word document: {e}")

# Main execution
if __name__ == "__main__":
    file_path = "./Amarin_CEA_Design_Specification.docx"

    document_text = extract_text_from_docx(file_path)

    if document_text:
        try:
            summarized_text = summarize_text(document_text, tokenizer)
            save_summary_to_doc(summarized_text, "./summary.docx")
        except Exception as e:
            print(f"Error during summarization: {e}")
    else:
        print("The document is empty or could not be read.")
