from transformers import AutoTokenizer, pipeline
import torch
from docx import Document

model_name = "facebook/bart-large-cnn"
summarizer = pipeline(task="summarization", model=model_name, device=0 if torch.cuda.is_available() else -1)
tokenizer = AutoTokenizer.from_pretrained(model_name)

def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        text = " ".join(paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip())
        return text
    except Exception as e:
        print(f"Error reading the .docx file: {e}")
        return ""

def split_text(text, tokenizer, max_tokens=1024):
    words = text.split()
    chunks = []
    current_chunk = []

    for word in words:
        current_chunk.append(word)
        if len(tokenizer(" ".join(current_chunk))["input_ids"]) >= max_tokens:
            chunks.append(" ".join(current_chunk[:-1]))
            current_chunk = [word]

    if current_chunk:
        chunks.append(" ".join(current_chunk))

    return chunks

def summarize_text(text, tokenizer, chunk_size=1024):
    chunks = split_text(text, tokenizer, max_tokens=chunk_size)
    summaries = []

    for i, chunk in enumerate(chunks):
        try:
            summary = summarizer(chunk, max_length=200, min_length=50, do_sample=False)
            summaries.append(summary[0]["summary_text"])
        except Exception as e:
            print(f"Error summarizing chunk {i}: {e}")

    return " ".join(summaries)

def generate_summary(input_file, output_file="./summary.docx"):
    document_text = extract_text_from_docx(input_file)

    if document_text:
        try:
            summarized_text = summarize_text(document_text, tokenizer)
            save_summary_to_doc(summarized_text, output_file)
            return summarized_text  # Return summarized text for use in the second script
        except Exception as e:
            print(f"Error during summarization: {e}")
            return None
    else:
        print("The document is empty or could not be read.")
        return None

def save_summary_to_doc(summary, output_path):
    try:
        doc = Document()
        doc.add_heading("Summarized Content", level=1)
        doc.add_paragraph(summary)
        doc.save(output_path)
        print(f"Summary saved successfully to {output_path}")
    except Exception as e:
        print(f"Error saving summary to Word document: {e}")
